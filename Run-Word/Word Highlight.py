import os
import os.path
import zipfile
from datetime import datetime
from re import sub
import re
import docx

import plac
import random
from pathlib import Path
import spacy
from tqdm import tqdm
import json
import io

#from nltk import sent_tokenize, word_tokenize
#import nltk
#from nltk.tag import SennaNERTagger





wordFile = r"C:\Users\zl0\Desktop\Testing NLP\Run-Word\IO\Input\ASAY0401.docx"
ttt = os.path.splitext(wordFile)[0]
ttr = os.path.basename(ttt)
key_Dict = {}
sourceFile = zipfile.ZipFile(wordFile, "r")

list = sourceFile.namelist()

publicDir = r"C:\Users\zl0\Desktop\Testing NLP\Run-Word\IO\Temp\Un-zip"
timestamp = datetime.now().isoformat()
timestamp = sub("[:.]", "_", timestamp)
dirname = ttr + timestamp
tempDir = os.path.join(publicDir, dirname)

os.mkdir(tempDir)

sourceFile.extractall(tempDir)

documentPath = os.path.join(tempDir, "word")

document = open(os.path.join(documentPath, "document.xml"), "a+", encoding='utf-8')
doc = open(os.path.join(documentPath, "document.xml"), "r+", encoding='utf-8')
Of = open(os.path.join(documentPath, "document1.xml"), 'w', encoding='utf-8')
for loop in doc.readlines():
    ent = re.sub(r"<w:r ([^>]+)>",r"<w:r>", loop)
    ent2 = re.sub(r"<w:r>", r"\n<w:r>", ent)
    ent1 = re.sub(r"</w:r>",r"</w:r>\n", ent2)
    ent1 = ent1.replace("<w:lastRenderedPageBreak/>","<peyar>")
    ent1 = ent1.replace("<w:t xml:space="'"preserve"'">","<w:t>")

    Of.write(ent1)
doc.close()
document.close()
Of.close()

os.rename(os.path.join(documentPath, "document.xml"), os.path.join(documentPath, "document2.xml"))

docxml = open(os.path.join(documentPath, "document.xml"), "w+")
of = open(wordFile , "rb")
docs = docx.Document(of)
n = 0
# docn = docx.Document()
doc = docx.Document(wordFile)
Of = open(os.path.join(documentPath, "document1.xml"), 'U')
nlp = spacy.load(r'C:\Users\zl0\Desktop\Testing NLP\T_Model\nl_model-0.0.0\nl_model\nl_model-0.0.0')
# nlp = spacy.load(r'C:\Users\zl0\Downloads\spaCy-master\spaCy-master\examples\training\Model\en_ANIPER-0.0.0\en_ANIPER\en_ANIPER-0.0.0')
for n in range(len(doc.paragraphs)):
    snt_t = doc.paragraphs[n].text
    doc_file = nlp(snt_t)
    print('doc_file:', doc_file)

    for word in doc_file.ents:
        print('word:', word.label_,'+', word.text)

        wlbl = word.label_
        wtxt = word.text
        # print(key_Dict.items())
        # print(bool(key_Dict))
#        if bool(key_Dict)=="":



        if wtxt not in key_Dict.keys():
            key = wtxt
            val = wlbl
            key_Dict[(key)] = val



        else:
            if wlbl not in key_Dict.items():
                key_Dict[(key)] = val

                        # create a new array in this slot
       # print(kwd[0])


print('key_Dict:', key_Dict)
with open(r"C:\Users\zl0\Desktop\Testing NLP\Run-Word\IO\Temp/" + ttr + ".txt", "a") as nf:
    for kWord in key_Dict.items():
        nf.write(kWord[0] + "|" + kWord[1] + "\n")
        # if wtxt not in kWord[0]:
            #print(kWord[0] + "|" + kWord[1])
nf.close()
for lines in Of.readlines():
    match = re.findall(r"<w:r><w:rPr>(.*)</w:rPr><w:t>", lines, re.M | re.I)
    if "<w:r>" in lines:

        for t in key_Dict.items():
            runs = re.findall(r"<w:t>(.*)</w:t>", lines, re.M | re.I)
            # print(runs)
            while t[0] in lines:
                for run in runs:
                # if t[0] !="de":
                    if match:
                        for mt in match:
                            if t[1] == "betrokkene":
                                if t[0] in run:
                                    lines = lines.replace(t[0],"</w:t></w:r><w:r><w:rPr>"+mt+"<w:highlight w:val="'"cyan"'"/></w:rPr><w:t>"+"&lt;ano as="'"'+t[1]+'"'"&gt;" + t[0] +"&lt;/ano&gt;"+ "</w:t></w:r><w:r><w:rPr>"+mt+"</w:rPr><w:t>")
                            if t[1] == "verzoeker":
                                if t[0] in run:
                                    lines = lines.replace(t[0],"</w:t></w:r><w:r><w:rPr>"+mt+"<w:highlight w:val="'"cyan"'"/></w:rPr><w:t>"+"&lt;ano as="'"'+t[1]+'"'"&gt;" + t[0] +"&lt;/ano&gt;"+  "</w:t></w:r><w:r><w:rPr>"+mt+"</w:rPr><w:t>")
                            if t[1] == "LOC":
                                if t[0] in run:
                                    lines = lines.replace(t[0],"</w:t></w:r><w:r><w:rPr>"+mt+"<w:highlight w:val="'"cyan"'"/></w:rPr><w:t>" +"&lt;ano as="'"'+t[1]+'"'"&gt;" + t[0] +"&lt;/ano&gt;"+ "</w:t></w:r><w:r><w:rPr>"+mt+"</w:rPr><w:t>")

                    else:
                        if t[1] == "betrokkene":
                            if t[0] in run:
                                lines = lines.replace(t[0], "</w:t></w:r><w:r><w:rPr><w:highlight w:val="'"cyan"'"/></w:rPr><w:t>"+"&lt;ano as="'"'+t[1]+'"'"&gt;" + t[0] +"&lt;/ano&gt;"+ "</w:t></w:r><w:r><w:t>")
                        if t[1] == "verzoeker":
                            if t[0] in run:
                                lines = lines.replace(t[0], "</w:t></w:r><w:r><w:rPr><w:highlight w:val="'"cyan"'"/></w:rPr><w:t>" +"&lt;ano as="'"'+t[1]+'"'"&gt;" + t[0] +"&lt;/ano&gt;"+  "</w:t></w:r><w:r><w:t>")
                        if t[1] == "LOC":
                            if t[0] in run:
                                lines = lines.replace(t[0], "</w:t></w:r><w:r><w:rPr><w:highlight w:val="'"cyan"'"/></w:rPr><w:t>" +"&lt;ano as="'"'+t[1]+'"'"&gt;" + t[0] +"&lt;/ano&gt;"+  "</w:t></w:r><w:r><w:t>")

                break
    lines = lines.replace("<w:t></w:t>","")
    lines = lines.replace("<w:r></w:r>", "")
    lines = lines.replace("<w:t> ", "<w:t xml:space="'"preserve"'"> ")
    lines = lines.replace(" </w:t>", "</w:t><w:t xml:space="'"preserve"'"> </w:t>")
    lines = lines.replace("<peyar>", "<w:lastRenderedPageBreak/>")
    # lines = lines.replace("<w:lastRen</w:t></w:r><w:r><w:rPr><w:highlight w:val="'"blue"'"/></w:rPr><w:t>de</w:t></w:r><w:r><w:t>redPageBreak/>","<w:lastRenderedPageBreak/>")
    O = open(os.path.join(documentPath, "document.xml"), 'a+')
    O.write(lines)
    O.close()


def zipdir(tempDir, ziph):
    length = len(tempDir)

    # ziph is zipfile handle
    for root, dirs, files in os.walk(tempDir):
        folder = root[length:] # path without "parent"
        for file in files:
            ziph.write(os.path.join(root, file), os.path.join(folder, file))

if __name__ == '__main__':
    zipf = zipfile.ZipFile(r'C:\Users\zl0\Desktop\Testing NLP\Run-Word\IO\Output/'+ttr+'.docx', 'w', zipfile.ZIP_DEFLATED)
    zipdir(tempDir, zipf)
    zipf.close()
