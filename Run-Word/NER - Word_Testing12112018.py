import docx
import os
import spacy

wordFile = r"C:\Users\zl0\Desktop\Testing NLP\Run-Word\IO\Input\ASVU0E01_notrck.docx"
ttt = os.path.splitext(wordFile)[0]
ttr = os.path.basename(ttt)
of = open(wordFile, "rb")
docs = docx.Document(of)
n = 0
docn = docx.Document()
nlp = spacy.load('nl')
doc = docx.Document(wordFile)
for n in range(len(docs.paragraphs)):
    snt_t = docs.paragraphs[n].text
    docFile = nlp(snt_t)
    # print(tokenized)
    for word in docFile.ents:
        wText = word.text
        wLabel = word.label_
        # if wLabel=="verdachte" or wLabel=="slachtoffer" or wLabel=="betrokkene" or wLabel=="appellant" or wLabel=="appellante" or wLabel=="naam" or wLabel=="eiser" or wLabel=="eiseres" or wLabel=="verbalisant" or wLabel=="belanghebbende" or wLabel=="de man" or wLabel=="de vader" or wLabel=="de vrouw" or wLabel=="de moeder" or wLabel=="het kind" or wLabel=="kind 1" or wLabel=="de zoon" or wLabel=="de dochter" or wLabel=="gedaagde" or wLabel=="geintimeerde" or wLabel=="gerekwireerde" or wLabel=="gerequireerde" or wLabel=="rekwirant" or wLabel=="requirant" or wLabel=="rekwirante" or wLabel=="requirante" or wLabel=="slachtoffer" or wLabel=="verzoeker" or wLabel=="verzoekster" or wLabel=="verweerder" or wLabel=="verweerster" or wLabel=="gemachtigde" or wLabel=="vertegenwoordiger" or wLabel=="directeur" or wLabel=="werknemer" or wLabel=="werkgever" or wLabel=="leidinggevende" or wLabel=="arts" or wLabel=="psycholoog" or wLabel=="psychiater" or wLabel=="veroordeelde" or wLabel=="getuige":
        #     Nlbl="Person"

        for p in doc.paragraphs:
            if wText in p.text:
                inline = p.runs

                for i in range(len(inline)):
                    if wText in inline[i].text:
                        text = inline[i].text.replace(wText, '<ano as="'+wLabel+'">'+wText+'</ano>')
                        inline[i].text = text

    doc.save(r"C:\Users\zl0\Desktop\Testing NLP\Run-Word\IO\Output/"+ttr+".docx")

    n = n + 1
